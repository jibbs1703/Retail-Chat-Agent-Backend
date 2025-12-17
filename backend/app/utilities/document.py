"""
Document generation utilities.
"""

import os
from io import BytesIO

import markdown
from docx import Document


def generate_docx_document(
    topic: str, paper_sections: dict[str, str], citations: list
) -> bytes:
    """Generate a DOCX document from paper content."""
    doc = Document()

    title = doc.add_heading(topic, 0)
    title.alignment = 1

    doc.add_paragraph()

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(paper_sections.get("abstract", ""))

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(paper_sections.get("introduction", ""))

    doc.add_heading("Literature Review", level=1)
    doc.add_paragraph(paper_sections.get("literature_review", ""))

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(paper_sections.get("discussion", ""))

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(paper_sections.get("conclusion", ""))

    doc.add_heading("References", level=1)
    for citation in citations:
        doc.add_paragraph(citation["formatted"], style="List Number")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def generate_markdown_document(
    topic: str, paper_sections: dict[str, str], citations: list
) -> str:
    """Generate a Markdown document from paper content."""
    md_content = f"# {topic}\n\n"

    md_content += "## Abstract\n\n"
    md_content += paper_sections.get("abstract", "") + "\n\n"

    md_content += "## Introduction\n\n"
    md_content += paper_sections.get("introduction", "") + "\n\n"

    md_content += "## Literature Review\n\n"
    md_content += paper_sections.get("literature_review", "") + "\n\n"

    md_content += "## Discussion\n\n"
    md_content += paper_sections.get("discussion", "") + "\n\n"

    md_content += "## Conclusion\n\n"
    md_content += paper_sections.get("conclusion", "") + "\n\n"

    md_content += "## References\n\n"
    for i, citation in enumerate(citations, 1):
        md_content += f"{i}. {citation['formatted']}\n"

    return md_content


def markdown_to_html(markdown_content: str) -> str:
    """Convert markdown to HTML."""
    return markdown.markdown(markdown_content)


def save_document(
    content: bytes, filename: str, directory: str = "storage/documents"
) -> str:
    """Save document to file system."""
    os.makedirs(directory, exist_ok=True)
    filepath = os.path.join(directory, filename)

    with open(filepath, "wb") as f:
        f.write(content)

    return filepath


def generate_research_package(
    topic: str, paper_sections: dict[str, str], citations: list
) -> dict[str, bytes]:
    """Generate a complete research package in multiple formats."""
    docx_content = generate_docx_document(topic, paper_sections, citations)

    md_content = generate_markdown_document(topic, paper_sections, citations)
    md_bytes = md_content.encode("utf-8")

    html_content = markdown_to_html(md_content)
    html_bytes = html_content.encode("utf-8")

    return {"docx": docx_content, "md": md_bytes, "html": html_bytes}
